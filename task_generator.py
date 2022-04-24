from PyQt5.QtWidgets import (QWidget, QSlider, QLineEdit, QLabel, QPushButton, QScrollArea,QApplication,
                             QHBoxLayout, QButtonGroup, QRadioButton, QVBoxLayout, QMainWindow, QDoubleSpinBox)
from PyQt5.QtCore import Qt, QSize
from PyQt5 import QtWidgets, uic
import sys
import random
import sqlite3
from PyQt5.QtWidgets import QMainWindow, QFileDialog
from errorr.error import Error
from errorr.unfaithful import Unfaithful
from docx import Document
from errorr.error_min_max import Error_min_max
from errorr.number_error import Number_error
from errorr.variables_error import Variables_error
from errorr.emptiness_error import Emptiness_error
from errorr.none_variable_error import None_variable_error
from errorr.none_formula_error import None_formula_error
from errorr.non_existent import Non_existent
from errorr.no_variable import No_variable
from PyQt5 import uic


class task_generator(QMainWindow):
    def __init__(self):
        super().__init__()
        uic.loadUi('testss.ui', self)
        self.con = sqlite3.connect('Tesks.db')
        self.cur = self.con.cursor()

        self.widget = QWidget()
        self.vbox = QVBoxLayout()
        self.list_variables = []
        self.save.clicked.connect(self.Save)
        self.Run.clicked.connect(self.Finding_variables)
        self.widget.setLayout(self.vbox)
        self.parmenial.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOn)
        self.parmenial.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.parmenial.setWidgetResizable(True)
        self.parmenial.setWidget(self.widget)

        self.save.hide()
        self.file.hide()
        self.file.clicked.connect(self.File)



    def Finding_variables(self):
        self.list_variables = []
        self.variables = []
        task_text = self.tasks_text.text()
        if task_text.count('&') == 0 or task_text.count('&') % 2 != 0:
            self.None_variable_error()
        else:
            self.lst_text_task = task_text.split()
            task_text = list(task_text)
            for i in range(task_text.count('&') // 2):
                a1 = task_text.index('&')
                del task_text[a1]
                a2 = task_text.index('&')
                del task_text[a2]
                self.variables.append(''.join(task_text[a1:a2]))
            for i in self.variables:
                if len(i) == 0:
                    self.Emptiness_error()
                    # введите название переменной
            for i in range(len(self.lst_text_task)):
                if self.lst_text_task[i].count('&') == 2:
                    a = list(self.lst_text_task[i])
                    del a[0]
                    del a[-1]
                    self.lst_text_task[i] = ''.join(a)
            self.change = self.lst_text_task.copy()
            self.save.show()

            for i in range(len(self.variables)):
                name_object = QLabel(self.variables[i])
                self.vbox.addWidget(name_object)
                min_object = QLabel('min')
                self.vbox.addWidget(min_object)
                self.min_p = QDoubleSpinBox()
                self.min_p.setMinimum(-10000)
                self.vbox.addWidget((self.min_p))
                max_object = QLabel('max')
                self.vbox.addWidget(max_object)
                self.max_p = QDoubleSpinBox()
                self.max_p.setMaximum(10000)
                self.vbox.addWidget((self.max_p))
                self.first_whole = QRadioButton('Целое')
                self.first_material = QRadioButton('Вещественное')
                lst_buttengroup = QButtonGroup()
                lst_buttengroup.addButton(self.first_whole)
                lst_buttengroup.addButton(self.first_material)

                self.vbox.addWidget(self.first_whole)
                self.vbox.addWidget(self.first_material)
                # добавление всех данных для переменной
                variables = [self.variables[i], self.min_p, self.max_p, lst_buttengroup, self.first_whole, self.first_material]
                self.list_variables.append({i: variables})

        self.save.show()

    def Formula(self):
        a = self.formula.text()
        if a == '':
            self.None_formula_error()
        self.lst_formula = list(a.split())
        self.check = ['*', '/', '-', '+']
        for i in self.variables:
            self.check.append(i)
            if i not in self.lst_formula:
                self.Non_existent()
        for i in self.lst_formula:
            if i not in self.check:
                if not i.isdigit():
                    self.Non_existent()
                else:
                    self.check.append(i)
        self.change_formula = self.lst_formula.copy()


    def Save(self):
        self.Formula()
        self.variable_data = []
        for i in range(len(self.variables)):
            a = float(self.list_variables[i][i][1].value())
            b = float(self.list_variables[i][i][2].value())
            self.variable_data.append(a)
            self.variable_data.append(b)
            if self.list_variables[i][i][5].isChecked():
                self.variable_data.append(1)
            else:
                self.variable_data.append(0)
        self.Savee()


    def Savee(self):
        a = self.formula.text()
        if a == '':
            self.None_formula_error()
        self.lst_formula = list(a.split())
        self.check = ['*', '/', '-', '+']
        for i in self.variables:
            self.check.append(i)
            if i not in self.lst_formula:
                self.Non_existent()
        for i in self.lst_formula:
            if i not in self.check:
                if not i.isdigit():
                    self.Non_existent()
                else:
                    self.check.append(i)
        self.change_formula = self.lst_formula.copy()
        self.cur.execute("""DELETE from tacks""")
        self.options = int(self.spinBox.text())

        a = self.formula.text()
        aa = set(self.lst_text_task)
        bb = set(self.lst_formula)
        if a == '':
            self.None_formula_error()
        elif len(aa & bb) != len(self.variables):
            self.Non_existent()
        else:
            for i in range(self.options):
                if len(self.variables) * 3 < len(self.variable_data):
                    self.Error()
                for j in range(len(self.variables)):
                    a = self.variable_data[j * 3]
                    b = self.variable_data[1 + (j * 3)]
                    if a > b or a == b:
                        self.Error_min_max()
                    if self.variable_data[2 + (j * 3)] == 0:
                        self.num1 = random.randint(int(a), int(b))
                    else:
                        self.num1 = round(random.uniform(a, b), 2)
                    self.lst_text_task[self.lst_text_task.index(self.variables[j])] = str(self.num1)
                    self.lst_formula[self.lst_formula.index(self.variables[j])] = self.num1

                self.lst_text_task = list(map(str, self.lst_text_task))
                self.lst_formula = list(map(str, self.lst_formula))
                lst = [(' '.join(self.lst_text_task), str(eval(' '.join(self.lst_formula))))]

                self.cur.executemany("INSERT INTO tacks VALUES (?, ?)", lst)
                self.con.commit()
                self.lst_text_task = self.change.copy()
                self.lst_formula = self.change_formula.copy()
                self.file.show()

    def File(self):
        a = self.cur.execute("""SELECT text FROM tacks""").fetchall()
        b = self.cur.execute("""SELECT ans FROM tacks""").fetchall()
        fname = QFileDialog.getSaveFileName()
        fname = fname[0]
        for i in range(self.options + 1):
            if i == 0:
                for j in range(self.options):
                    if j == 0:
                        self.document = Document()
                        self.document.add_heading(f'Ответы', level=1)
                        self.document.add_paragraph(f'{j + 1} - {b[j]}')
                        self.document.save(fname)
                    else:
                        self.document = Document(fname)
                        self.document.add_paragraph(f'{j + 1} - {b[j]}')
                        self.document.save(fname)
            else:
                self.document = Document(fname)
                self.document.add_paragraph()
                self.document.add_page_break()
                self.document.add_heading(f'Вариант {i}', level=1)
                self.document.add_paragraph(a[i - 1])
                self.document.save(fname)


    def Unfaithful(self):
        self.fourth = Unfaithful()
        self.fourth.show()

    def Error(self):
        self.fourth = Error()
        self.fourth.show()

    def Back(self):
        self.hide()

    def Error_min_max(self):
        self.fourth = Error_min_max()
        self.fourth.show()

    def Number_error(self):
        self.fourth = Number_error()
        self.fourth.show()

    def Variables_error(self):
        self.fourth = Variables_error()
        self.fourth.show()

    def Emptiness_error(self):
        self.fourth = Emptiness_error()
        self.fourth.show()

    def None_variable_error(self):
        self.fourth = None_variable_error()
        self.fourth.show()

    def None_formula_error(self):
        self.fourth = None_formula_error()
        self.fourth.show()

    def Non_existent(self):
        self.fourth = Non_existent()
        self.fourth.show()

    def No_variable(self):
        self.fourth = No_variable()
        self.fourth.show()